"""DocxReportAdapter — generates audit-ready Word documents using python-docx.

Implements domain.ports.report_ports.ReportGeneratorPort while keeping all
infrastructure concerns (file I/O, styling) outside the domain layer.
"""

from __future__ import annotations

import logging
import os
from datetime import datetime, timezone
from typing import List, TYPE_CHECKING

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from domain.entities.reconciliation import CertificateOfAccuracy

if TYPE_CHECKING:
    from application.use_cases.migration_pipeline import MigrationResult

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------
_GREEN = RGBColor(0x22, 0x8B, 0x22)   # ForestGreen
_RED = RGBColor(0xCC, 0x00, 0x00)     # Dark red
_GREY = RGBColor(0x66, 0x66, 0x66)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_HEADER_BG = "2F5496"                  # dark-blue table header


def _set_cell_shading(cell, hex_colour: str) -> None:
    """Apply a solid background colour to a table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_colour)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


class DocxReportAdapter:
    """Generates audit-ready .docx reports from domain entities."""

    # ------------------------------------------------------------------
    # Certificate of Accuracy report
    # ------------------------------------------------------------------

    async def generate_certificate_report(
        self, certificate: CertificateOfAccuracy, output_path: str
    ) -> str:
        """Create a professional Word document for a single certificate."""
        doc = Document()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10)

        # ---- Title ----
        title = doc.add_heading("Certificate of Accuracy", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ---- Status badge ----
        is_certified = certificate.status == "CERTIFIED"
        badge_text = certificate.status
        badge = doc.add_heading(badge_text, level=1)
        badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in badge.runs:
            run.font.color.rgb = _GREEN if is_certified else _RED

        # ---- Header details ----
        doc.add_paragraph("")  # spacer
        details = [
            ("Certificate ID", certificate.certificate_id),
            ("Module", certificate.module),
            ("Period", f"{certificate.period.period_name} "
                       f"({certificate.period.period_year}-{certificate.period.period_num:02d})"),
            ("Org ID", str(certificate.context.org_id)),
            ("Issued At", certificate.issued_at.strftime("%Y-%m-%d %H:%M:%S UTC")),
            ("Status", certificate.status),
        ]
        for label, value in details:
            p = doc.add_paragraph()
            run_label = p.add_run(f"{label}: ")
            run_label.bold = True
            p.add_run(value)

        # ---- Reconciliation checks table ----
        doc.add_heading("Reconciliation Checks", level=2)

        headers = [
            "Check Type", "Source", "Target",
            "Source Value", "Target Value",
            "Variance", "Tolerance", "Status",
        ]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Style header row
        for idx, text in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = _WHITE
                    run.font.size = Pt(9)
            _set_cell_shading(cell, _HEADER_BG)

        # Data rows
        for check in certificate.result.checks:
            row = table.add_row()
            passed = check.is_within_tolerance
            cells = row.cells
            cells[0].text = check.check_type
            cells[1].text = check.source_label
            cells[2].text = check.target_label
            cells[3].text = str(check.source_value)
            cells[4].text = str(check.target_value)
            cells[5].text = str(check.variance)
            cells[6].text = str(check.tolerance)
            status_text = "PASS" if passed else "FAIL"
            cells[7].text = status_text
            # Colour the status cell
            for paragraph in cells[7].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = _GREEN if passed else _RED
                    run.font.bold = True

        # ---- Signature block ----
        doc.add_paragraph("")
        sig = doc.add_paragraph()
        sig_run = sig.add_run(f"Issued by: {certificate.issuer}")
        sig_run.italic = True

        # ---- Footer with timestamp ----
        doc.add_paragraph("")
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
        run_footer = footer.add_run(f"Report generated on {ts}")
        run_footer.font.size = Pt(8)
        run_footer.font.color.rgb = _GREY

        # ---- Save ----
        filepath = self._resolve_path(output_path, f"{certificate.certificate_id}.docx")
        doc.save(filepath)
        logger.info("Certificate report saved to %s", filepath)
        return filepath

    # ------------------------------------------------------------------
    # Migration Summary report
    # ------------------------------------------------------------------

    async def generate_migration_summary(
        self, results: List, output_path: str
    ) -> str:
        """Create a summary Word document for multiple migration runs."""
        doc = Document()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10)

        # ---- Title ----
        title = doc.add_heading("Migration Summary Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ---- Summary table ----
        doc.add_heading("Migration Runs", level=2)

        headers = [
            "Module", "Period", "Rows Extracted", "Rows Loaded",
            "Reconciliation", "Certificate", "Duration (s)",
        ]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for idx, text in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = _WHITE
                    run.font.size = Pt(9)
            _set_cell_shading(cell, _HEADER_BG)

        for r in results:
            row = table.add_row()
            cells = row.cells
            cells[0].text = r.module
            cells[1].text = (
                f"{r.period.period_name} "
                f"({r.period.period_year}-{r.period.period_num:02d})"
            )
            cells[2].text = str(r.rows_extracted)
            cells[3].text = str(r.rows_loaded)
            recon_text = "PASS" if r.reconciliation_passed else "FAIL"
            cells[4].text = recon_text
            for paragraph in cells[4].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = _GREEN if r.reconciliation_passed else _RED
                    run.font.bold = True
            cells[5].text = r.certificate_id or "N/A"
            cells[6].text = str(r.duration_seconds)

        # ---- Overall stats ----
        doc.add_heading("Overall Statistics", level=2)

        total_modules = len(results)
        total_rows = sum(r.rows_extracted for r in results)
        passed = sum(1 for r in results if r.reconciliation_passed)
        pass_rate = (passed / total_modules * 100) if total_modules else 0

        stats = [
            ("Total Modules", str(total_modules)),
            ("Total Rows Extracted", str(total_rows)),
            ("Pass Rate", f"{pass_rate:.1f}%"),
        ]
        for label, value in stats:
            p = doc.add_paragraph()
            run_label = p.add_run(f"{label}: ")
            run_label.bold = True
            p.add_run(value)

        # ---- Footer ----
        doc.add_paragraph("")
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
        run_footer = footer.add_run(f"Report generated on {ts}")
        run_footer.font.size = Pt(8)
        run_footer.font.color.rgb = _GREY

        filepath = self._resolve_path(output_path, "migration_summary.docx")
        doc.save(filepath)
        logger.info("Migration summary report saved to %s", filepath)
        return filepath

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _resolve_path(output_path: str, default_name: str) -> str:
        """If *output_path* is a directory, append *default_name*; otherwise use as-is."""
        if os.path.isdir(output_path):
            return os.path.join(output_path, default_name)
        return output_path
